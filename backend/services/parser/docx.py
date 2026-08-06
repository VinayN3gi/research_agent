import docx
from models import Document
from services.parser.dispatcher import BaseParser
import os

class DocxParser(BaseParser):
    async def parse(self, file_path: str, metadata: dict = None) -> Document:
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        filename = os.path.basename(file_path)
        return Document(
            id=file_path,
            title=metadata.get("title", filename),
            source_type="docx",
            text=text.strip(),
            metadata=metadata or {}
        )
